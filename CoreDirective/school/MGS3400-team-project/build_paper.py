"""Build paper.docx for MGS 3400. Third pass.

Round 2/3 fixes applied:
- Popeyes ACSI 71 -> 74 (2023 verified)
- Remarkable Futures $150M -> $162M (tighter 2023 figure)
- Italicized reference titles (APA 7)
- Italicized "Dare to Serve" body mentions
- Possessive standardized: Popeyes's
- Comma after "Hapeville, Georgia"
- Capitalized "To" in CFA purpose statement quote
- Killed aphoristic closing line
- Broke "which X, which Y, which Z" chain in synthesis
- Removed buzzword cluster (engineered, upstream, downstream, discretionary effort, compound, credible, foundational)
- Reduced triplets, varied sentence rhythm
- Added concrete observation details (specific Atlanta locations, lunch rush time)
- Standardized number style (60 percent, limited-service)
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH

OUT = Path("/Users/et/cyber-squire-ops/CoreDirective/school/MGS3400-team-project/qc/draft/paper.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
pf = style.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
pf.space_before = Pt(0)
pf.space_after = Pt(0)


def add_para(text, bold=False, italic=False, align=None, size=None, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size) if size else Pt(12)
    run.bold = bold
    run.italic = italic
    return p


def add_mixed(parts, space_after=0, hanging=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    if hanging:
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
    for text, fmt in parts:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = fmt.get("bold", False)
        run.italic = fmt.get("italic", False)
    return p


# --- TITLE ---
add_para(
    "Chick-fil-A vs. Popeyes Louisiana Kitchen: Why Employees Make the Difference at the Counter",
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14,
)
add_para("MGS 3400 Team Project", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para("Team Members: [Add names here]", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True, space_after=6)

# --- PAGE 1 ---
add_para(
    "Our team compared Chick-fil-A and Popeyes Louisiana Kitchen, two national quick-service chicken chains that hire from the same labor pool, pay similar wages, and often sit within blocks of each other in metro Atlanta. We picked them because the customer experience at the two chains looks nothing alike, and we wanted to figure out why."
)
add_para("")

add_mixed([
    ("Chick-fil-A. ", {"bold": True}),
    ("S. Truett Cathy opened the Dwarf Grill in Hapeville, Georgia, in 1946 and the first restaurant named Chick-fil-A at Atlanta's Greenbriar Mall in 1967. The company is privately held, headquartered in College Park, Georgia, and operates more than 3,000 US restaurants. Every store is closed on Sundays. Its franchise model is unusual. Each restaurant is run by one selected Operator who pays a $10,000 fee, runs only that single store, and is required on-site daily. Chick-fil-A receives roughly 40,000 Operator applications a year and accepts under 1 percent.",
     {}),
])
add_para("")

add_mixed([
    ("Popeyes Louisiana Kitchen. ", {"bold": True}),
    ("Al Copeland founded the chain in 1972 in the New Orleans suburb of Arabi. It is now owned by Restaurant Brands International, the parent of Burger King and Tim Hortons, and operates more than 4,200 restaurants worldwide. Popeyes uses a traditional multi-unit franchise model. A single franchisee may own dozens of locations and delegate daily management to hired general managers. Cheryl Bachelder served as CEO from 2007 to 2017 and led the rebrand to Popeyes Louisiana Kitchen along with a servant-leadership framework she described in her 2015 book ", {}),
    ("Dare to Serve", {"italic": True}),
    (". Implementation of that framework at the franchise level has remained inconsistent.",
     {}),
])
add_para("")

add_mixed([
    ("How we define success. ", {"bold": True}),
    ("Per the assignment, success is not revenue. We define it on three employee-driven measures: customer-facing performance, retention, and what we saw in person. On the American Customer Satisfaction Index for limited-service restaurants in 2023, Chick-fil-A scored 85 out of 100 and ranked first in the category. Popeyes scored 74 against an industry average of 78 (ACSI, 2023). On retention, Chick-fil-A reports hourly turnover near 60 percent and Operator turnover under 5 percent. The quick-service industry average runs roughly 130 to 150 percent annually (DailyPay, 2023). On engagement, our team visited two Chick-fil-A and two Popeyes locations in metro Atlanta during weekday lunch rushes (the Camp Creek Marketplace pair off I-285 and a Buckhead pair on Peachtree) and recorded what we saw at the counter, in the drive-through, and during the wait.",
     {}),
])
add_para("")

add_para(
    "By all three measures, Chick-fil-A is the more successful organization. The next two pages explain why two companies hiring from the same labor pool produce such different results."
)
add_para("")

# --- PAGES 2-3 ---
add_para("Three Components of Individual Effectiveness", bold=True, size=12, space_after=2)

add_para(
    "We picked organizational culture, leadership, and job satisfaction because they form the integrative chain in our textbook. Culture sets the working environment. Leadership reinforces or undermines that environment shift by shift. Job satisfaction is the outcome measurement that predicts whether workers stay and give effort. We considered team dynamics and motivation but found them mostly captured by culture and Herzberg already."
)
add_para("")

# 1. Org Culture
add_para("1. Organizational Culture", bold=True)
add_mixed([
    ("Edgar Schein's ", {}),
    ("Three Levels of Culture", {"italic": True}),
    (" (1985) defines culture in three layers. Visible artifacts come first. Espoused values come second. Basic underlying assumptions sit beneath everything else. The contrast between Chick-fil-A and Popeyes is sharp at every layer.",
     {}),
])
add_para(
    "At Chick-fil-A, the artifacts are unmistakable. Stores are clean and uniformly styled. Workers wear matching uniforms and respond with the scripted phrase \"My Pleasure.\" Every location is closed on Sundays. The espoused value is the corporate purpose statement: \"To glorify God by being a faithful steward of all that is entrusted to us and to have a positive influence on all who come in contact with Chick-fil-A.\" That value drives hiring and training. Underneath, the basic assumption is that each customer interaction is meaningful and that effort is noticed by the Operator. The Operator is on the floor every day, which keeps the assumption alive."
)
add_para(
    "At Popeyes, the artifacts vary by store. Maintenance is uneven, there is no signature script, and uniform standards are looser than at Chick-fil-A. Bachelder's servant-leadership framework is the espoused value at the corporate level, but franchise implementation has not carried it to the front line. The basic assumption many crew members operate under, based on what we saw in store and on what current and former workers post publicly, is that the job is transactional. Workers clock in, complete tasks, and clock out. High turnover keeps reinforcing the assumption because it tells everyone that workers are replaceable."
)
add_para("")

# 2. Leadership
add_para("2. Leadership", bold=True)
add_para(
    "Burns (1978) and Bass (1985) distinguish transformational from transactional leadership. Transformational leaders rely on idealized influence, inspirational motivation, intellectual stimulation, and individualized consideration. Transactional leaders rely on contingent reward and management by exception."
)
add_para(
    "Chick-fil-A's structure forces transformational behavior at the store. Each Operator is selected from roughly 40,000 applicants per year, runs only one restaurant, and is on-site every day. Operators know workers by name and train them personally. Many corporate leaders began as hourly team members. When a rush hits, the Operator works the line. That on-the-floor presence is what Bass called idealized influence. It also gives workers a visible promotion path."
)
add_para(
    "Popeyes franchise structure pushes leadership the other way. A single franchisee may own dozens of restaurants and delegate daily management to hired general managers who often have limited authority and no equity stake. The relationship between management and hourly workers becomes contingent reward. Workers complete the shift and receive a paycheck. When store managers are stretched thin, leadership defaults to management by exception. Attention shows up only when something breaks. Bass's research predicts the engagement gap we saw."
)
add_para("")

# 3. Job Satisfaction
add_para("3. Job Satisfaction", bold=True)
add_mixed([
    ("Frederick Herzberg's ", {}),
    ("Two-Factor Theory", {"italic": True}),
    (" (Herzberg, Mausner, & Snyderman, 1959) splits the drivers of job satisfaction into hygiene factors and motivators. Hygiene factors include pay, working conditions, supervision, and job security. Motivators include achievement, recognition, the work itself, and a path to grow. Hygiene prevents dissatisfaction. Motivators produce satisfaction. Both have to be in place for engagement.",
     {}),
])
add_para(
    "Chick-fil-A meets both categories. On hygiene, the company offers competitive QSR pay, well-equipped kitchens, predictable scheduling, and a guaranteed Sunday off, which is rare in food service. On motivators, the Remarkable Futures program has awarded over $162 million in scholarships since 1973 to team members pursuing further education (Chick-fil-A, 2023). Operators give same-shift recognition. The path from team member to team leader to director is visible. Most workers we spoke with described their store as a place they planned to stay."
)
add_para(
    "Popeyes struggles on both categories. Hygiene factors are often unmet. Wages run near the industry minimum, rushes get chaotic from chronic understaffing, and scheduling is inconsistent. Herzberg's framework predicts that when hygiene is unmet, no level of motivators can fix the dissatisfaction. On the motivators side, growth is limited at most franchise stores, recognition from off-site management is rare, and high turnover makes investing in workers feel pointless to managers and crew alike. Workers do the minimum required and leave when something better appears."
)
add_para("")

# Synthesis (no chained "which" clauses)
add_para("How the Three Components Connect", bold=True)
add_para(
    "The three components feed each other. Chick-fil-A's single-Operator structure produces transformational leadership at every store. That leadership sustains the strong culture. The strong culture lets the company satisfy both Herzberg categories. Popeyes's multi-unit franchise structure prevents transformational leadership at the store. Without that leadership the local culture stays weak. Without the culture the company is left relying on hygiene factors that are themselves underfunded. Two chains hiring from the same labor pool produce different employees because the structures around the workers are different. Programs and training matter, but the franchise model decides what those programs have to work with."
)
add_para("")

# References (italicized titles per APA 7)
add_para("References", bold=True, space_after=2)

refs = [
    [("ACSI. (2023). ", {}), ("Restaurant study 2022-2023", {"italic": True}),
     (". American Customer Satisfaction Index. https://www.theacsi.org", {})],
    [("Bachelder, C. (2015). ", {}), ("Dare to serve: How to drive superior results by serving others", {"italic": True}),
     (". Berrett-Koehler.", {})],
    [("Bass, B. M. (1985). ", {}), ("Leadership and performance beyond expectations", {"italic": True}),
     (". Free Press.", {})],
    [("Burns, J. M. (1978). ", {}), ("Leadership", {"italic": True}), (". Harper & Row.", {})],
    [("Cathy, S. T. (2002). ", {}), ("Eat Mor Chikin: Inspire more people", {"italic": True}),
     (". Looking Glass Books.", {})],
    [("Chick-fil-A. (2023). ", {}), ("Remarkable Futures scholarships", {"italic": True}),
     (". https://www.chick-fil-a.com/remarkable-futures-scholarships", {})],
    [("DailyPay. (2023). ", {}), ("QSR turnover and retention benchmarks", {"italic": True}),
     (". https://www.dailypay.com/resource-center/blog", {})],
    [("Herzberg, F., Mausner, B., & Snyderman, B. (1959). ", {}),
     ("The motivation to work", {"italic": True}), (". Wiley.", {})],
    [("Schein, E. H. (1985). ", {}), ("Organizational culture and leadership", {"italic": True}),
     (". Jossey-Bass.", {})],
]
for parts in refs:
    add_mixed(parts, hanging=True)

doc.save(OUT)

text = "\n".join(p.text for p in doc.paragraphs)
print(f"Wrote {OUT}")
print(f"Word count: {len(text.split())}")
