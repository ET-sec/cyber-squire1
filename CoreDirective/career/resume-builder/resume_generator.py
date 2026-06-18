#!/usr/bin/env python3
"""Resume generator: JSON data -> formatted .docx resume.

Reads a base resume_data.json and optional variant overlay,
produces a single-column .docx with Calibri typography,
bold metrics, hyperlinked contact info, and section borders.
"""

import argparse
import copy
import json
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, RGBColor
from docx.opc.constants import RELATIONSHIP_TYPE as RT


# ---------------------------------------------------------------------------
# Data loading
# ---------------------------------------------------------------------------

def _deep_merge(base: dict, overlay: dict) -> dict:
    """Recursively merge *overlay* into *base*, returning a new dict."""
    merged = copy.deepcopy(base)
    for key, value in overlay.items():
        if key in merged and isinstance(merged[key], dict) and isinstance(value, dict):
            merged[key] = _deep_merge(merged[key], value)
        else:
            merged[key] = copy.deepcopy(value)
    return merged


def load_resume_data(base_path: str, variant_path: str | None = None) -> dict:
    """Load base JSON. If *variant_path* given, deep-merge overrides.

    Variant JSON may override: title, contact_line, certifications,
    skills_override, and provide bullet_order mapping (job key -> ordered
    list of bullet keys).
    """
    with open(base_path, "r", encoding="utf-8") as fh:
        data = json.load(fh)

    if variant_path is None:
        return data

    with open(variant_path, "r", encoding="utf-8") as fh:
        variant = json.load(fh)

    # Top-level scalar overrides
    for field in ("title", "headline", "contact_line", "certifications"):
        if field in variant:
            data[field] = variant[field]

    # Skills override replaces entire skills list
    if "skills_override" in variant:
        data["skills"] = variant["skills_override"]

    # Education override replaces entire education list
    if "education_override" in variant:
        data["education"] = variant["education_override"]

    # Contact override replaces entire contact block
    if "contact_override" in variant:
        data["contact"] = variant["contact_override"]

    # Per-job overrides (job_title, company, location, dates)
    if "job_overrides" in variant:
        job_map = {job["key"]: job for job in data.get("experience", []) if "key" in job}
        for job_key, overrides in variant["job_overrides"].items():
            if job_key in job_map:
                for field in ("job_title", "company", "location", "dates"):
                    if field in overrides:
                        job_map[job_key][field] = overrides[field]

    # bullet_order: map variant job keys to experience entries by their `key`
    if "bullet_order" in variant:
        job_map = {job["key"]: job for job in data.get("experience", []) if "key" in job}
        for job_key, ordered_bullets in variant["bullet_order"].items():
            if job_key in job_map:
                job_map[job_key]["bullet_order"] = ordered_bullets

    # bullet_overrides: per-job per-bullet text rewrites (variant-only edits)
    if "bullet_overrides" in variant:
        job_map = {job["key"]: job for job in data.get("experience", []) if "key" in job}
        for job_key, bullet_map in variant["bullet_overrides"].items():
            if job_key in job_map:
                for bkey, btext in bullet_map.items():
                    job_map[job_key].setdefault("bullets", {})[bkey] = btext

    return data


# ---------------------------------------------------------------------------
# Hyperlink helper
# ---------------------------------------------------------------------------

def add_hyperlink(paragraph, url: str, text: str):
    """Insert a clickable hyperlink run into *paragraph*.

    Creates an external relationship, builds the w:hyperlink XML element,
    and styles the run with blue underlined text.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Blue colour
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "0000FF")
    rPr.append(color_el)

    # Underline
    u_el = OxmlElement("w:u")
    u_el.set(qn("w:val"), "single")
    rPr.append(u_el)

    # Font
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)

    # Size (match contact line: 10pt = 20 half-points)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "20")
    rPr.append(szCs)

    new_run.append(rPr)

    t_el = OxmlElement("w:t")
    t_el.set(qn("xml:space"), "preserve")
    t_el.text = text
    new_run.append(t_el)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ---------------------------------------------------------------------------
# Section header with bottom border
# ---------------------------------------------------------------------------

def add_section_header(doc, title: str):
    """Add an all-caps section header with a thin bottom border."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.3

    run = para.add_run(title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    # Bottom border via XML
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")       # 4 = ~0.5pt line
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Bold-metric bullet helper
# ---------------------------------------------------------------------------

BOLD_RE = re.compile(r'\*\*(.+?)\*\*')


def add_bullet_with_bold_metrics(paragraph, text: str):
    """Clear paragraph runs and re-add with **markers** bolded."""
    for run in paragraph.runs:
        run._r.getparent().remove(run._r)

    parts = BOLD_RE.split(text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        if idx % 2 == 1:  # bold match
            run.bold = True


# ---------------------------------------------------------------------------
# Job entry
# ---------------------------------------------------------------------------

def add_job_entry(doc, job: dict):
    """Add a single experience block: title | company, location | dates, bullets."""
    # Title line
    para_title = doc.add_paragraph()
    para_title.paragraph_format.space_before = Pt(2)
    para_title.paragraph_format.space_after = Pt(0)
    para_title.paragraph_format.line_spacing = 1.3
    run_title = para_title.add_run(job.get("job_title", ""))
    run_title.bold = True
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(10)

    run_sep = para_title.add_run("  |  ")
    run_sep.font.name = "Calibri"
    run_sep.font.size = Pt(10)

    run_company = para_title.add_run(job.get("company", ""))
    run_company.font.name = "Calibri"
    run_company.font.size = Pt(10)

    # Location / dates line
    para_loc = doc.add_paragraph()
    para_loc.paragraph_format.space_before = Pt(0)
    para_loc.paragraph_format.space_after = Pt(2)
    para_loc.paragraph_format.line_spacing = 1.3
    loc_text = f'{job.get("location", "")}  |  {job.get("dates", "")}'
    run_loc = para_loc.add_run(loc_text)
    run_loc.italic = True
    run_loc.font.name = "Calibri"
    run_loc.font.size = Pt(10)

    # Bullets -- order controlled by bullet_order list
    bullets_dict = job.get("bullets", {})
    bullet_order = job.get("bullet_order", list(bullets_dict.keys()))

    for bkey in bullet_order:
        btext = bullets_dict.get(bkey)
        if btext is None:
            continue
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after = Pt(0)
        bp.paragraph_format.line_spacing = 1.3
        add_bullet_with_bold_metrics(bp, btext)


# ---------------------------------------------------------------------------
# Skills category
# ---------------------------------------------------------------------------

def add_skill_category(doc, category: dict):
    """Add a skills line: bold label followed by regular items."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.line_spacing = 1.3
    label = category.get("label", "")
    items = category.get("items", "")

    run_label = para.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.name = "Calibri"
    run_label.font.size = Pt(10)

    run_items = para.add_run(items)
    run_items.font.name = "Calibri"
    run_items.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Projects
# ---------------------------------------------------------------------------

def _add_project_entry(doc, proj: dict):
    """Add a single project bullet: bold name + hyperlink, description after."""
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(2)

    name = proj.get("name", "")
    url = proj.get("url", "")
    desc = proj.get("description", "")

    # Bold name run
    run_name = para.add_run(name)
    run_name.bold = True
    run_name.font.name = "Calibri"
    run_name.font.size = Pt(10)

    # URL as hyperlink
    if url:
        sep = para.add_run("  |  ")
        sep.font.name = "Calibri"
        sep.font.size = Pt(10)
        add_hyperlink(para, url, url)

    # Description
    if desc:
        sep2 = para.add_run(" — ")
        sep2.font.name = "Calibri"
        sep2.font.size = Pt(10)

        parts = BOLD_RE.split(desc)
        for idx, part in enumerate(parts):
            if not part:
                continue
            run = para.add_run(part)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            if idx % 2 == 1:
                run.bold = True


# ---------------------------------------------------------------------------
# Education
# ---------------------------------------------------------------------------

def add_education(doc, education: dict):
    """Add education block: school, degrees, honors."""
    para_school = doc.add_paragraph()
    para_school.paragraph_format.space_before = Pt(1)
    para_school.paragraph_format.space_after = Pt(0)
    para_school.paragraph_format.line_spacing = 1.3
    run_school = para_school.add_run(education.get("school", ""))
    run_school.bold = True
    run_school.font.name = "Calibri"
    run_school.font.size = Pt(10)

    if "degrees" in education:
        para_deg = doc.add_paragraph()
        para_deg.paragraph_format.space_before = Pt(0)
        para_deg.paragraph_format.space_after = Pt(0)
        para_deg.paragraph_format.line_spacing = 1.3
        run_deg = para_deg.add_run(education["degrees"])
        run_deg.font.name = "Calibri"
        run_deg.font.size = Pt(10)

    if "additional" in education:
        para_add = doc.add_paragraph()
        para_add.paragraph_format.space_before = Pt(0)
        para_add.paragraph_format.space_after = Pt(0)
        para_add.paragraph_format.line_spacing = 1.3
        run_add = para_add.add_run(education["additional"])
        run_add.font.name = "Calibri"
        run_add.font.size = Pt(10)

    if "honors" in education:
        para_hon = doc.add_paragraph()
        para_hon.paragraph_format.space_before = Pt(0)
        para_hon.paragraph_format.space_after = Pt(1)
        para_hon.paragraph_format.line_spacing = 1.3
        run_hon = para_hon.add_run(education["honors"])
        run_hon.italic = True
        run_hon.font.name = "Calibri"
        run_hon.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Contact line with hyperlinks
# ---------------------------------------------------------------------------

def _build_contact_line(doc, data: dict):
    """Build left-aligned contact paragraph with plain text and hyperlinks."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.3

    contact = data.get("contact", {})
    parts = contact.get("parts", [])
    links = contact.get("links", {})

    # Simple string fallback
    if not parts and "contact_line" in data:
        run = para.add_run(data["contact_line"])
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        return

    for i, part in enumerate(parts):
        if i > 0:
            sep = para.add_run(" | ")
            sep.font.name = "Calibri"
            sep.font.size = Pt(10)

        if part in links:
            add_hyperlink(para, links[part], part)
        else:
            run = para.add_run(part)
            run.font.name = "Calibri"
            run.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Main document creation
# ---------------------------------------------------------------------------

def create_resume(data: dict, output_path: str):
    """Generate a .docx resume from structured *data* and save to *output_path*."""
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.top_margin = Inches(0.25)
    section.bottom_margin = Inches(0.25)
    section.gutter = Inches(0)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)

    # ── Name ──
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_para.paragraph_format.space_after = Pt(2)
    name_para.paragraph_format.line_spacing = 1.3
    name_run = name_para.add_run(data.get("name", ""))
    name_run.bold = True
    name_run.font.name = "Calibri"
    name_run.font.size = Pt(16)

    # ── Headline (title line under name, above contact) ──
    headline = data.get("headline")
    if headline:
        hp = doc.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after = Pt(2)
        hr = hp.add_run(headline)
        hr.font.name = "Calibri"
        hr.font.size = Pt(11)
        hr.italic = False

    # ── Contact line ──
    _build_contact_line(doc, data)

    # ── CERTIFICATIONS ──
    if "certifications" in data:
        add_section_header(doc, "CERTIFICATIONS")
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(1)
        cp.paragraph_format.space_after = Pt(1)
        cr = cp.add_run(data["certifications"])
        cr.bold = True
        cr.font.name = "Calibri"
        cr.font.size = Pt(10)

    if "certifications_pursuing" in data:
        cp2 = doc.add_paragraph()
        cp2.paragraph_format.space_before = Pt(0)
        cp2.paragraph_format.space_after = Pt(1)
        cr2 = cp2.add_run(data["certifications_pursuing"])
        cr2.font.name = "Calibri"
        cr2.font.size = Pt(10)
        cr2.italic = True

    # ── EXPERIENCE ──
    experience = data.get("experience", [])
    if experience:
        add_section_header(doc, "EXPERIENCE")
        for job in experience:
            add_job_entry(doc, job)

    # ── TECHNICAL SKILLS ──
    skills = data.get("skills", [])
    if skills:
        add_section_header(doc, "TECHNICAL SKILLS")
        for cat in skills:
            add_skill_category(doc, cat)

    # ── PROJECTS ──
    projects = data.get("projects", [])
    if projects:
        add_section_header(doc, "PROJECTS")
        for proj in projects:
            _add_project_entry(doc, proj)

    # ── EDUCATION ──
    education_list = data.get("education", [])
    if education_list:
        add_section_header(doc, "EDUCATION")
        for edu in education_list:
            add_education(doc, edu)

    # ── Save ──
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    print(f"Resume saved to {output_path}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Generate a formatted .docx resume from JSON data."
    )
    parser.add_argument(
        "--base", required=True,
        help="Path to base resume_data.json"
    )
    parser.add_argument(
        "--variant", default=None,
        help="Optional path to variant overlay JSON"
    )
    parser.add_argument(
        "--output", default="output/Emmanuel_Tigoue_AI_Security_Engineer.docx",
        help="Output .docx path (default: output/Emmanuel_Tigoue_AI_Security_Engineer.docx)"
    )
    args = parser.parse_args()

    data = load_resume_data(args.base, args.variant)
    create_resume(data, args.output)


if __name__ == "__main__":
    main()
