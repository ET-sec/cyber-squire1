#!/usr/bin/env python3
"""BS / fraud detection scoring for resume .docx files.

Scores 10 dimensions of potential resume fraud.  Each detected signal
adds 1 point.  Lower score = cleaner resume.

Target: 4/10 or lower (some signals are acceptable).
"""

import argparse
import json
import re
import sys

from docx import Document


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_full_text(docx_path: str) -> str:
    """Return all paragraph text from a .docx, newline-separated."""
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs)


def extract_bullets(docx_path: str) -> list[str]:
    """Return bullet texts."""
    doc = Document(docx_path)
    bullets = []
    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower()
        text = para.text.strip()
        if text and ("list" in style_name or "bullet" in style_name):
            bullets.append(text)
    return bullets


# ---------------------------------------------------------------------------
# 10 BS signal checks
# ---------------------------------------------------------------------------

VAGUE_WORDS = ["significant", "substantial", "numerous", "various",
               "multiple", "several", "considerable"]
INFLATED_TITLES = ["senior", "principal", "director", "vp",
                   "vice president", "chief"]
BUZZWORDS = ["synergy", "paradigm", "ecosystem", "holistic",
             "transformative", "innovative"]
MONTH_YEAR_RE = re.compile(
    r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\.?\s+(\d{4})',
    re.IGNORECASE,
)
MONTH_MAP = {
    "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
    "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
}


def sig_vague_metrics(bullets: list[str]) -> tuple[int, str]:
    """Any bullet with vague quantity words instead of numbers?"""
    lower_bullets = [b.lower() for b in bullets]
    found = [w for w in VAGUE_WORDS
             if any(w in b for b in lower_bullets)]
    return (1 if found else 0,
            f"Vague words: {', '.join(found)}" if found else "None found")


def sig_title_inflation(text: str) -> tuple[int, str]:
    """Contains Senior/Principal/Director/VP/Chief for < 5 yrs experience?"""
    lower = text.lower()
    # Estimate years from date ranges
    dates = MONTH_YEAR_RE.findall(text)
    years_set = {int(y) for _, y in dates}
    if years_set:
        span = max(years_set) - min(years_set)
    else:
        span = 0

    inflated = [t for t in INFLATED_TITLES if t in lower]
    if inflated and span < 5:
        return 1, f"Titles {inflated} with ~{span}yr span"
    return 0, f"No inflation (span={span}yr)"


def sig_scale_mismatch(text: str) -> tuple[int, str]:
    """Claims 'enterprise' scale with non-Fortune-500 company?"""
    lower = text.lower()
    enterprise_words = ["enterprise", "fortune", "global scale",
                        "multinational"]
    has_enterprise = any(w in lower for w in enterprise_words)
    # Simple heuristic: if enterprise claim exists, flag it for review
    # Real checker would cross-reference company names
    if has_enterprise:
        return 1, "Enterprise-scale language detected"
    return 0, "No enterprise claims"


def sig_impossible_metrics(bullets: list[str]) -> tuple[int, str]:
    """Any percentage > 99% or cost savings > $10M for non-exec?"""
    for b in bullets:
        # Check for > 99% improvements
        pcts = re.findall(r'(\d+)%', b)
        for p in pcts:
            if int(p) > 99:
                return 1, f"Found {p}% claim"

        # Check for > $10M savings
        money = re.findall(r'\$(\d[\d,]*\.?\d*)\s*([MBT])?', b, re.IGNORECASE)
        for amount, suffix in money:
            val = float(amount.replace(",", ""))
            if suffix and suffix.upper() == "M" and val > 10:
                return 1, f"Found ${amount}{suffix} claim"
            if suffix and suffix.upper() in ("B", "T"):
                return 1, f"Found ${amount}{suffix} claim"

    return 0, "Metrics within plausible range"


def sig_tool_namedropping(bullets: list[str]) -> tuple[int, str]:
    """5+ tool names in a single bullet with no use-case context?"""
    # Heuristic: bullets shorter than 15 words with 5+ capitalized/acronym words
    tool_pattern = re.compile(r'\b[A-Z][A-Za-z]*(?:\s[A-Z][A-Za-z]*)*\b')
    for b in bullets:
        words = b.split()
        if len(words) < 15:
            tools = tool_pattern.findall(b)
            # Filter out common non-tool capitalized words
            tools = [t for t in tools if len(t) > 1
                     and t not in ("I", "The", "A", "An", "In", "On", "At",
                                   "For", "To", "By", "Of", "And", "Or")]
            if len(tools) >= 5:
                return 1, f"Bullet with {len(tools)} tool names: '{b[:50]}...'"
    return 0, "No tool name-dropping detected"


def sig_buzzword_density(text: str) -> tuple[int, str]:
    """Count of buzzwords > 3?"""
    lower = text.lower()
    count = sum(lower.count(bw) for bw in BUZZWORDS)
    return (1 if count > 3 else 0,
            f"{count} buzzwords found")


def sig_identical_structures(bullets: list[str]) -> tuple[int, str]:
    """4+ bullets with identical grammatical structure?"""
    if len(bullets) < 4:
        return 0, "Too few bullets"

    # Simple heuristic: first word + rough structure
    patterns = []
    for b in bullets:
        words = b.split()
        if len(words) >= 3:
            # Pattern = first word + whether 2nd word is capitalized + word count bucket
            pat = (words[0].lower(),
                   words[1][0].isupper() if len(words) > 1 else False,
                   len(words) // 5)
            patterns.append(pat)

    from collections import Counter
    counts = Counter(patterns)
    max_count = max(counts.values()) if counts else 0
    if max_count >= 4:
        return 1, f"{max_count} bullets with same structure"
    return 0, f"Max {max_count} similar structures"


def sig_education_padding(text: str) -> tuple[int, str]:
    """Listing 'Relevant Coursework' or 3+ awards/honors?"""
    lower = text.lower()
    has_coursework = "relevant coursework" in lower or "coursework:" in lower
    awards = re.findall(r'(award|honor|dean|magna|summa|cum laude)', lower)
    if has_coursework or len(awards) > 3:
        return 1, f"Coursework={has_coursework}, awards={len(awards)}"
    return 0, "Minimal education padding"


def sig_cert_overclaiming(text: str) -> tuple[int, str]:
    """Listing expired or unverifiable certs?"""
    lower = text.lower()
    if "expired" in lower:
        return 1, "Expired cert mentioned"
    # Check for obscure/unverifiable cert names -- very basic heuristic
    return 0, "No expired certs detected"


def sig_employment_overlap(text: str) -> tuple[int, str]:
    """Two full-time roles with overlapping date ranges?"""
    # Find all date ranges (start - end) per line
    lines = text.split("\n")
    ranges = []
    for line in lines:
        dates = MONTH_YEAR_RE.findall(line)
        if len(dates) >= 2:
            # Parse first and last date in the line
            start_month = MONTH_MAP.get(dates[0][0][:3].lower(), 1)
            start_year = int(dates[0][1])
            end_month = MONTH_MAP.get(dates[-1][0][:3].lower(), 12)
            end_year = int(dates[-1][1])
            start_val = start_year * 12 + start_month
            end_val = end_year * 12 + end_month
            ranges.append((start_val, end_val))
        elif "present" in line.lower() and len(dates) >= 1:
            start_month = MONTH_MAP.get(dates[0][0][:3].lower(), 1)
            start_year = int(dates[0][1])
            start_val = start_year * 12 + start_month
            end_val = 2030 * 12  # far future
            ranges.append((start_val, end_val))

    # Check for overlaps (excluding "Present")
    for i in range(len(ranges)):
        for j in range(i + 1, len(ranges)):
            s1, e1 = ranges[i]
            s2, e2 = ranges[j]
            if s1 < e2 and s2 < e1:
                return 1, "Overlapping date ranges detected"

    return 0, "No overlapping dates"


# ---------------------------------------------------------------------------
# Scoring
# ---------------------------------------------------------------------------

def score_bs_detection(docx_path: str) -> dict:
    """Run all 10 signal checks and return structured results."""
    text = extract_full_text(docx_path)
    bullets = extract_bullets(docx_path)

    signals = {
        "vague_metrics": sig_vague_metrics(bullets),
        "title_inflation": sig_title_inflation(text),
        "scale_mismatch": sig_scale_mismatch(text),
        "impossible_metrics": sig_impossible_metrics(bullets),
        "tool_namedropping": sig_tool_namedropping(bullets),
        "buzzword_density": sig_buzzword_density(text),
        "identical_structures": sig_identical_structures(bullets),
        "education_padding": sig_education_padding(text),
        "cert_overclaiming": sig_cert_overclaiming(text),
        "employment_overlap": sig_employment_overlap(text),
    }

    bs_score = sum(score for score, _ in signals.values())

    detected = [name for name, (score, _) in signals.items() if score == 1]
    clean = [name for name, (score, _) in signals.items() if score == 0]

    breakdown = {
        name: {"detected": score == 1, "detail": detail}
        for name, (score, detail) in signals.items()
    }

    return {
        "bs_score": bs_score,
        "max": 10,
        "signals_detected": detected,
        "clean_signals": clean,
        "breakdown": breakdown,
        "pass": bs_score <= 4,
    }


# ---------------------------------------------------------------------------
# Output
# ---------------------------------------------------------------------------

TARGET_SCORE = 4


def print_report(result: dict):
    """Print human-readable BS detection report."""
    print("=" * 60)
    print("BS / FRAUD DETECTION REPORT")
    print("=" * 60)
    print(f"BS Score: {result['bs_score']}/{result['max']} (lower is better)")
    status = "PASS" if result["pass"] else "FAIL"
    print(f"Status: {status} (target <= {TARGET_SCORE}/10)")
    print()

    print("Signals Detected:")
    for name, info in result["breakdown"].items():
        label = name.replace("_", " ").title()
        mark = "[!]" if info["detected"] else "[ ]"
        print(f"  {mark} {label}: {info['detail']}")

    print()
    print(f"Clean: {len(result['clean_signals'])}/10")
    print(f"Flagged: {len(result['signals_detected'])}/10")
    print("=" * 60)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Score a resume .docx for BS/fraud signals."
    )
    parser.add_argument("docx_path", help="Path to the resume .docx file")
    parser.add_argument(
        "--output", choices=["human", "json"], default="human",
        help="Output format (default: human)"
    )
    args = parser.parse_args()

    result = score_bs_detection(args.docx_path)

    if args.output == "json":
        print(json.dumps(result, indent=2))
    else:
        print_report(result)

    sys.exit(0 if result["pass"] else 1)


if __name__ == "__main__":
    main()
