#!/usr/bin/env python3
"""AI detection heuristic scoring for resume .docx files.

Scores bullet text against anti-AI checklist heuristics.
Higher human_score = more human-like writing.
ai_detection_score = inverse, on 0-10 scale (lower is better).

Target: ai_detection_score 3/10 or lower.
"""

import argparse
import json
import math
import re
import sys

from docx import Document


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_bullets(docx_path: str) -> list[str]:
    """Return list of bullet/list-paragraph texts from a .docx."""
    doc = Document(docx_path)
    bullets = []
    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower()
        text = para.text.strip()
        if not text:
            continue
        # Capture bullets from List Bullet style or any list-like paragraph
        if "list" in style_name or "bullet" in style_name:
            bullets.append(text)
    # Fallback: if no styled bullets found, treat lines starting with common
    # bullet chars as bullets
    if not bullets:
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and (text.startswith("-") or text.startswith("*")
                         or text.startswith("\u2022")):
                bullets.append(text.lstrip("-* \u2022").strip())
    return bullets


# ---------------------------------------------------------------------------
# Heuristic checks
# ---------------------------------------------------------------------------

BANNED_VERBS = ["spearheaded", "orchestrated", "leveraged", "utilized", "facilitated"]
OPENING_VERBS_RE = re.compile(r'^(\w+)')


def _first_word(text: str) -> str:
    m = OPENING_VERBS_RE.match(text)
    return m.group(1).lower() if m else ""


def check_bullet_opening_variety(bullets: list[str]) -> tuple[int, str]:
    """0-15 pts. Deduct 5 for each group of 3+ consecutive same-opening bullets."""
    if len(bullets) < 3:
        return 15, "Too few bullets to evaluate"

    score = 15
    openings = [_first_word(b) for b in bullets]
    streak = 1
    deductions = 0
    for i in range(1, len(openings)):
        if openings[i] == openings[i - 1] and openings[i]:
            streak += 1
        else:
            if streak >= 3:
                deductions += 1
            streak = 1
    if streak >= 3:
        deductions += 1

    score = max(0, score - deductions * 5)
    return score, f"{deductions} groups of 3+ same-opening bullets"


def check_sentence_length_variance(bullets: list[str]) -> tuple[int, str]:
    """0-15 pts based on stddev of bullet word counts."""
    if len(bullets) < 2:
        return 15, "Too few bullets"

    lengths = [len(b.split()) for b in bullets]
    mean = sum(lengths) / len(lengths)
    variance = sum((x - mean) ** 2 for x in lengths) / len(lengths)
    stddev = math.sqrt(variance)

    if stddev > 5:
        score = 15
    elif stddev > 3:
        score = 10
    elif stddev > 2:
        score = 5
    else:
        score = 0

    return score, f"stddev={stddev:.2f}"


def check_informal_narrative(bullets: list[str]) -> tuple[int, str]:
    """0-15 pts. Check for narrative patterns showing human writing."""
    detections = 0
    for b in bullets:
        # Fragment without standard verb opening
        words = b.split()
        if words and words[0][0].islower():
            detections += 1
            continue
        # First person pronouns
        if re.search(r'\b(I|my|we|our)\b', b):
            detections += 1
            continue
        # Parenthetical context asides
        if re.search(r'\(.*?\)', b) and len(b.split()) < 25:
            detections += 1

    score = min(15, detections * 5)
    return score, f"{detections} narrative patterns found"


def check_semicolons(bullets: list[str]) -> tuple[int, str]:
    """0-10 pts. 10 if zero semicolons, 0 if any."""
    text = " ".join(bullets)
    has_semicolons = ";" in text
    score = 0 if has_semicolons else 10
    return score, "semicolons found" if has_semicolons else "no semicolons"


def check_hyphens(bullets: list[str]) -> tuple[int, str]:
    """0-10 pts. 10 if no hyphens/em-dashes in bullet text."""
    text = " ".join(bullets)
    # Check for hyphens used as dashes (not in compound words)
    has_dashes = bool(re.search(r' [-\u2013\u2014] ', text))
    score = 0 if has_dashes else 10
    return score, "dashes found" if has_dashes else "no dashes"


def check_number_specificity(bullets: list[str]) -> tuple[int, str]:
    """0-15 pts. Non-round numbers are more human; round numbers less so."""
    text = " ".join(bullets)
    numbers = re.findall(r'\b(\d+)\b', text)

    score = 15
    for num_str in numbers:
        num = int(num_str)
        if num == 0:
            continue
        # Round = divisible by 5 or 10
        if num % 10 == 0 or num % 5 == 0:
            score -= 3
        else:
            score += 1

    score = max(0, min(15, score))
    return score, f"{len(numbers)} numbers found"


def check_banned_verbs(bullets: list[str]) -> tuple[int, str]:
    """0-10 pts. Deduct 2 per banned verb occurrence."""
    text = " ".join(bullets).lower()
    score = 10
    found = []
    for verb in BANNED_VERBS:
        count = text.count(verb)
        if count > 0:
            score -= 2 * count
            found.append(f"{verb}({count})")

    score = max(0, score)
    detail = ", ".join(found) if found else "none found"
    return score, detail


def check_metric_density(bullets: list[str]) -> tuple[int, str]:
    """0-10 pts. Score based on % of bullets containing metrics."""
    if not bullets:
        return 0, "no bullets"

    metric_re = re.compile(r'\$?\d[\d,]*\.?\d*[KMB]?%?(?:\+)?')
    with_metrics = sum(1 for b in bullets if metric_re.search(b))
    density = with_metrics / len(bullets) * 100

    if 55 <= density <= 75:
        score = 10
    elif 40 <= density < 55 or 75 < density <= 85:
        score = 5
    else:
        score = 0

    return score, f"{density:.0f}% ({with_metrics}/{len(bullets)})"


# ---------------------------------------------------------------------------
# Scoring
# ---------------------------------------------------------------------------

def score_ai_detection(docx_path: str) -> dict:
    """Run all 8 heuristic checks and return structured results."""
    bullets = extract_bullets(docx_path)

    checks = {
        "bullet_opening_variety": check_bullet_opening_variety(bullets),
        "sentence_length_variance": check_sentence_length_variance(bullets),
        "informal_narrative": check_informal_narrative(bullets),
        "semicolon_check": check_semicolons(bullets),
        "hyphen_check": check_hyphens(bullets),
        "number_specificity": check_number_specificity(bullets),
        "banned_verb_check": check_banned_verbs(bullets),
        "metric_density": check_metric_density(bullets),
    }

    human_score = sum(score for score, _ in checks.values())
    human_score = max(0, min(100, human_score))

    ai_detection_score = round((100 - human_score) / 10, 1)
    ai_detection_score = min(10.0, max(0.0, ai_detection_score))

    breakdown = {
        name: {"score": score, "max": max_pts, "detail": detail}
        for (name, (score, detail)), max_pts in zip(
            checks.items(),
            [15, 15, 15, 10, 10, 15, 10, 10]
        )
    }

    return {
        "ai_detection_score": ai_detection_score,
        "human_score": human_score,
        "bullets_analyzed": len(bullets),
        "breakdown": breakdown,
        "pass": ai_detection_score <= 3.0,
    }


# ---------------------------------------------------------------------------
# Output
# ---------------------------------------------------------------------------

TARGET_SCORE = 3.0


def print_report(result: dict):
    """Print human-readable AI detection report."""
    print("=" * 60)
    print("AI DETECTION HEURISTIC REPORT")
    print("=" * 60)
    print(f"AI Detection Score: {result['ai_detection_score']}/10 "
          f"(lower is better)")
    print(f"Human Score: {result['human_score']}/100")
    print(f"Bullets Analyzed: {result['bullets_analyzed']}")
    status = "PASS" if result["pass"] else "FAIL"
    print(f"Status: {status} (target <= {TARGET_SCORE})")
    print()

    for name, info in result["breakdown"].items():
        label = name.replace("_", " ").title()
        print(f"  {label}: {info['score']}/{info['max']} -- {info['detail']}")

    print("=" * 60)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Score a resume .docx for AI detection risk."
    )
    parser.add_argument("docx_path", help="Path to the resume .docx file")
    parser.add_argument(
        "--output", choices=["human", "json"], default="human",
        help="Output format (default: human)"
    )
    args = parser.parse_args()

    result = score_ai_detection(args.docx_path)

    if args.output == "json":
        print(json.dumps(result, indent=2))
    else:
        print_report(result)

    sys.exit(0 if result["pass"] else 1)


if __name__ == "__main__":
    main()
