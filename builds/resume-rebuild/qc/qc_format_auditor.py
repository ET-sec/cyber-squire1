#!/usr/bin/env python3
"""Formatting & structure audit for resume .docx files.

Checks 10 structural quality signals in the .docx XML and document
properties.  Each signal is 0 or 1 point (total /10).

Target: 7/10 or higher.
"""

import argparse
import json
import os
import re
import sys
import zipfile

from docx import Document
from docx.shared import Inches


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _get_docx_xml(docx_path: str) -> str:
    """Extract word/document.xml from the .docx (which is a ZIP)."""
    with zipfile.ZipFile(docx_path, "r") as z:
        return z.read("word/document.xml").decode("utf-8")


def _get_rels_xml(docx_path: str) -> str:
    """Extract word/_rels/document.xml.rels from the .docx."""
    with zipfile.ZipFile(docx_path, "r") as z:
        return z.read("word/_rels/document.xml.rels").decode("utf-8")


# ---------------------------------------------------------------------------
# 10 format checks
# ---------------------------------------------------------------------------

REQUIRED_HEADERS = ["CERTIFICATIONS", "EXPERIENCE", "TECHNICAL SKILLS", "EDUCATION"]
BANNED_HEADERS = ["SUMMARY", "OBJECTIVE", "PROFILE", "ABOUT"]


def check_section_headers(docx_path: str) -> tuple[int, str]:
    """Exactly 4 required section headers, no SUMMARY/OBJECTIVE."""
    doc = Document(docx_path)
    found_required = []
    found_banned = []

    for para in doc.paragraphs:
        text = para.text.strip().upper()
        if text in REQUIRED_HEADERS:
            found_required.append(text)
        if text in BANNED_HEADERS:
            found_banned.append(text)

    missing = set(REQUIRED_HEADERS) - set(found_required)
    if not missing and not found_banned:
        return 1, f"All 4 headers present, no banned headers"
    detail_parts = []
    if missing:
        detail_parts.append(f"missing: {', '.join(missing)}")
    if found_banned:
        detail_parts.append(f"banned: {', '.join(found_banned)}")
    return 0, "; ".join(detail_parts)


def check_contact_hyperlinks(docx_path: str) -> tuple[int, str]:
    """Contact line has w:hyperlink XML elements for LinkedIn, GitHub, Portfolio."""
    xml = _get_docx_xml(docx_path)
    # Count w:hyperlink elements
    hyperlink_count = xml.count("<w:hyperlink")
    # Check for the three expected link texts
    expected = ["LinkedIn", "GitHub", "Portfolio"]
    found = [label for label in expected if label in xml]

    if len(found) >= 3 and hyperlink_count >= 3:
        return 1, f"{hyperlink_count} hyperlinks, all 3 labels present"
    return 0, f"{hyperlink_count} hyperlinks, labels found: {found}"


def check_hyperlinks_clickable(docx_path: str) -> tuple[int, str]:
    """All hyperlinks have r:id relationships pointing to external URLs."""
    xml = _get_docx_xml(docx_path)
    rels_xml = _get_rels_xml(docx_path)

    # Extract r:id values from w:hyperlink elements
    hyperlink_rids = re.findall(r'<w:hyperlink[^>]*r:id="([^"]+)"', xml)
    if not hyperlink_rids:
        return 0, "No hyperlink r:id references found"

    # Check that each r:id maps to an external URL in rels
    missing_rels = []
    for rid in hyperlink_rids:
        # Look for Relationship with this Id pointing to external URL
        pattern = f'Id="{rid}"[^>]*Target="(https?://[^"]+)"'
        if not re.search(pattern, rels_xml):
            missing_rels.append(rid)

    if not missing_rels:
        return 1, f"All {len(hyperlink_rids)} hyperlinks have valid external targets"
    return 0, f"{len(missing_rels)} hyperlinks missing external target: {missing_rels}"


def check_name_centered_bold(docx_path: str) -> tuple[int, str]:
    """First non-empty paragraph is centered and bold (the name)."""
    doc = Document(docx_path)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Check alignment
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        is_centered = para.alignment == WD_ALIGN_PARAGRAPH.CENTER
        # Check bold
        is_bold = all(run.bold for run in para.runs if run.text.strip())
        if is_centered and is_bold:
            return 1, f"Name '{text}' is centered and bold"
        detail_parts = []
        if not is_centered:
            detail_parts.append("not centered")
        if not is_bold:
            detail_parts.append("not bold")
        return 0, f"Name '{text}': {', '.join(detail_parts)}"

    return 0, "No non-empty paragraphs found"


def check_margins(docx_path: str) -> tuple[int, str]:
    """Margins are 0.5 inches all around."""
    doc = Document(docx_path)
    section = doc.sections[0]

    target = Inches(0.5)
    # Allow 1% tolerance for floating point
    tolerance = target * 0.02

    margins = {
        "left": section.left_margin,
        "right": section.right_margin,
        "top": section.top_margin,
        "bottom": section.bottom_margin,
    }

    off = []
    for name, val in margins.items():
        if val is None or abs(val - target) > tolerance:
            actual_inches = val / 914400 if val else 0  # EMU to inches
            off.append(f"{name}={actual_inches:.2f}in")

    if not off:
        return 1, "All margins 0.5in"
    return 0, f"Off margins: {', '.join(off)}"


def check_font_consistency(docx_path: str) -> tuple[int, str]:
    """All runs use Calibri font."""
    doc = Document(docx_path)
    non_calibri = set()

    for para in doc.paragraphs:
        for run in para.runs:
            if not run.text.strip():
                continue
            font_name = run.font.name
            if font_name and font_name != "Calibri":
                non_calibri.add(font_name)

    # Also check hyperlink runs via XML since python-docx may not expose them
    # through normal run iteration
    if not non_calibri:
        return 1, "All runs use Calibri"
    return 0, f"Non-Calibri fonts found: {', '.join(non_calibri)}"


def check_no_orphaned_blanks(docx_path: str) -> tuple[int, str]:
    """No consecutive blank paragraphs (orphaned whitespace)."""
    doc = Document(docx_path)
    prev_blank = False
    orphan_count = 0

    for para in doc.paragraphs:
        is_blank = not para.text.strip()
        if is_blank and prev_blank:
            orphan_count += 1
        prev_blank = is_blank

    if orphan_count == 0:
        return 1, "No orphaned blank paragraphs"
    return 0, f"{orphan_count} consecutive blank paragraph pairs"


def check_bullet_style(docx_path: str) -> tuple[int, str]:
    """Bullet points use List Bullet style."""
    doc = Document(docx_path)
    list_style_count = 0
    non_list_bullet_count = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name or "").lower()

        # If it looks like a bullet (starts with common bullet chars)
        if text.startswith("\u2022") or text.startswith("-") or text.startswith("*"):
            if "list" not in style_name and "bullet" not in style_name:
                non_list_bullet_count += 1

        if "list" in style_name or "bullet" in style_name:
            list_style_count += 1

    if list_style_count > 0 and non_list_bullet_count == 0:
        return 1, f"{list_style_count} bullets using List Bullet style"
    if list_style_count == 0:
        return 0, "No List Bullet styled paragraphs found"
    return 0, f"{non_list_bullet_count} bullets not using List Bullet style"


def check_file_size(docx_path: str) -> tuple[int, str]:
    """File size between 10KB and 500KB."""
    size = os.path.getsize(docx_path)
    size_kb = size / 1024

    if 10 <= size_kb <= 500:
        return 1, f"{size_kb:.1f}KB (within 10-500KB range)"
    return 0, f"{size_kb:.1f}KB (outside 10-500KB range)"


def check_page_count_estimate(docx_path: str) -> tuple[int, str]:
    """Text length suggests 1-2 pages (roughly 2000-6000 chars)."""
    doc = Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    char_count = len(text)

    # A typical single-column resume page is ~2500-3500 chars
    # 1 page: 1500-4000 chars, 2 pages: 4000-7000 chars
    if 1500 <= char_count <= 7000:
        estimated_pages = 1 if char_count < 4000 else 2
        return 1, f"{char_count} chars (~{estimated_pages} page{'s' if estimated_pages > 1 else ''})"
    return 0, f"{char_count} chars (outside 1-2 page range)"


# ---------------------------------------------------------------------------
# Scoring
# ---------------------------------------------------------------------------

def score_format_audit(docx_path: str) -> dict:
    """Run all 10 format checks and return structured results."""
    checks = {
        "section_headers": check_section_headers(docx_path),
        "contact_hyperlinks": check_contact_hyperlinks(docx_path),
        "hyperlinks_clickable": check_hyperlinks_clickable(docx_path),
        "name_centered_bold": check_name_centered_bold(docx_path),
        "margins_half_inch": check_margins(docx_path),
        "font_consistency": check_font_consistency(docx_path),
        "no_orphaned_blanks": check_no_orphaned_blanks(docx_path),
        "bullet_style": check_bullet_style(docx_path),
        "file_size": check_file_size(docx_path),
        "page_count_estimate": check_page_count_estimate(docx_path),
    }

    total = sum(score for score, _ in checks.values())

    breakdown = {
        name: {"score": score, "detail": detail, "pass": score == 1}
        for name, (score, detail) in checks.items()
    }

    return {
        "format_score": total,
        "max": 10,
        "breakdown": breakdown,
        "pass": total >= 7,
    }


# ---------------------------------------------------------------------------
# Output
# ---------------------------------------------------------------------------

TARGET_SCORE = 7


def print_report(result: dict):
    """Print human-readable format audit report."""
    print("=" * 60)
    print("FORMAT & STRUCTURE AUDIT REPORT")
    print("=" * 60)
    print(f"Format Score: {result['format_score']}/{result['max']}")
    status = "PASS" if result["pass"] else "FAIL"
    print(f"Status: {status} (target >= {TARGET_SCORE}/10)")
    print()

    for name, info in result["breakdown"].items():
        label = name.replace("_", " ").title()
        mark = "[X]" if info["pass"] else "[ ]"
        print(f"  {mark} {label}: {info['detail']}")

    print("=" * 60)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Audit a resume .docx for formatting and structural quality."
    )
    parser.add_argument("docx_path", help="Path to the resume .docx file")
    parser.add_argument(
        "--output", choices=["human", "json"], default="human",
        help="Output format (default: human)"
    )
    args = parser.parse_args()

    result = score_format_audit(args.docx_path)

    if args.output == "json":
        print(json.dumps(result, indent=2))
    else:
        print_report(result)

    sys.exit(0 if result["pass"] else 1)


if __name__ == "__main__":
    main()
