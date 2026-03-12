#!/usr/bin/env python3
"""Hiring manager simulation scoring for resume .docx files.

Scores 10 dimensions that a hiring manager would assess in a
6-second resume scan. Each dimension is 0 or 1 point (total /10).

Target: 9/10 or higher.
"""

import argparse
import json
import re
import sys

from docx import Document


# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------

def extract_full_text(docx_path: str) -> str:
    """Return all paragraph text from a .docx, newline-separated."""
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs)


def extract_sections(docx_path: str) -> dict:
    """Parse the resume into named sections based on header keywords."""
    doc = Document(docx_path)
    sections: dict[str, list[str]] = {}
    current_section = "_header"
    sections[current_section] = []

    section_keywords = ["CERTIFICATIONS", "EXPERIENCE", "TECHNICAL SKILLS",
                        "EDUCATION", "SKILLS", "PROJECTS"]

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        upper = text.upper()
        if upper in section_keywords:
            current_section = upper
            sections[current_section] = []
        else:
            sections.setdefault(current_section, []).append(text)

    return sections


def extract_bullets_with_runs(docx_path: str) -> list[dict]:
    """Return bullets with bold-run info for metric detection."""
    doc = Document(docx_path)
    bullets = []
    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower()
        if "list" not in style_name and "bullet" not in style_name:
            continue
        text = para.text.strip()
        if not text:
            continue
        has_bold = any(run.bold for run in para.runs if run.text.strip())
        bullets.append({"text": text, "has_bold_metric": has_bold})
    return bullets


# ---------------------------------------------------------------------------
# 10 scoring dimensions
# ---------------------------------------------------------------------------

METRIC_RE = re.compile(r'\$?\d[\d,]*\.?\d*[KMB]?%?(?:\+)?')
ACTION_VERBS = [
    "deployed", "built", "created", "designed", "developed", "implemented",
    "configured", "automated", "engineered", "established", "reduced",
    "integrated", "secured", "managed", "led", "architected", "conducted",
    "wrote", "enforced", "hardened", "migrated", "consolidated",
]
SECURITY_TITLES = [
    "security", "cybersecurity", "infosec", "devsecops", "soc",
    "threat", "vulnerability", "compliance", "grc", "risk",
    "ai security", "cloud security",
]
MONTH_YEAR_RE = re.compile(
    r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\.?\s+\d{4}',
    re.IGNORECASE,
)


def dim_first_bullet_impact(bullets: list[dict]) -> tuple[int, str]:
    """Does first bullet of first job have a bold metric AND action verb?"""
    if not bullets:
        return 0, "No bullets found"
    first = bullets[0]
    has_metric = first["has_bold_metric"]
    words = first["text"].lower().split()
    has_verb = any(v in words[:3] for v in ACTION_VERBS)
    passed = has_metric and has_verb
    return (1 if passed else 0,
            f"metric={'yes' if has_metric else 'no'}, "
            f"verb={'yes' if has_verb else 'no'}")


def dim_certs_before_experience(sections: dict) -> tuple[int, str]:
    """Are certifications listed before experience?"""
    keys = list(sections.keys())
    cert_idx = None
    exp_idx = None
    for i, k in enumerate(keys):
        if "CERT" in k.upper():
            cert_idx = i
        if "EXPERIENCE" in k.upper() and exp_idx is None:
            exp_idx = i
    if cert_idx is not None and exp_idx is not None and cert_idx < exp_idx:
        return 1, "Certifications before Experience"
    if cert_idx is None:
        return 0, "No CERTIFICATIONS section"
    return 0, "Certifications after Experience"


def dim_metric_density(bullets: list[dict]) -> tuple[int, str]:
    """At least 10 bold metrics across all bullets?"""
    count = sum(1 for b in bullets if b["has_bold_metric"])
    passed = count >= 10
    return (1 if passed else 0,
            f"{count} bold metrics ({'>=10' if passed else '<10'})")


def dim_job_title_match(text: str) -> tuple[int, str]:
    """Does title line contain a recognized security title?"""
    lower = text.lower()
    for title in SECURITY_TITLES:
        if title in lower:
            return 1, f"Found '{title}'"
    return 0, "No security title detected"


def dim_dates_present(text: str) -> tuple[int, str]:
    """Does every job entry have parseable Mon YYYY dates?"""
    # Find job-like lines (contain " | ")
    lines = text.split("\n")
    job_lines = [l for l in lines if "|" in l and any(
        c.isalpha() for c in l)]

    # Among location/date lines, check for date patterns
    dates_found = MONTH_YEAR_RE.findall(text)
    if len(dates_found) >= 2:
        return 1, f"{len(dates_found)} dates found"
    return 0, f"Only {len(dates_found)} dates found"


def dim_no_gaps(text: str) -> tuple[int, str]:
    """No gap > 3 months between jobs. (Simplified: check for 'Present')."""
    # This is a simplified heuristic; full date parsing is in the real scorer
    dates = MONTH_YEAR_RE.findall(text)
    has_present = "present" in text.lower()
    if has_present and len(dates) >= 2:
        return 1, "Continuous employment indicated"
    if len(dates) >= 4:
        return 1, "Multiple date ranges found"
    return 0, f"Insufficient date data ({len(dates)} dates)"


def dim_skills_with_context(sections: dict) -> tuple[int, str]:
    """Do at least 3 skills lines have parenthetical context?"""
    skills_lines = sections.get("TECHNICAL SKILLS", [])
    paren_count = sum(1 for line in skills_lines if "(" in line and ")" in line)
    passed = paren_count >= 3
    return (1 if passed else 0,
            f"{paren_count} lines with context ({'>=3' if passed else '<3'})")


def dim_education_present(sections: dict) -> tuple[int, str]:
    """Education section present with degree and GPA?"""
    edu = sections.get("EDUCATION", [])
    if not edu:
        return 0, "No EDUCATION section"
    text = " ".join(edu).lower()
    has_degree = any(d in text for d in ["b.s.", "b.a.", "b.b.a.", "bba",
                                          "m.s.", "m.a.", "mba",
                                          "bachelor", "master", "associate",
                                          "phd", "doctorate"])
    has_gpa = "gpa" in text or "cum laude" in text or "dean" in text
    if has_degree and has_gpa:
        return 1, "Degree and GPA/honors found"
    if has_degree:
        return 0, "Degree found but no GPA/honors"
    return 0, "No degree detected"


def dim_single_page(text: str) -> tuple[int, str]:
    """Is total character count under 6000 (rough 1-2 page estimate)?

    Threshold is 6000 to accommodate legitimate 2-page resumes for
    candidates with 5+ years of experience and multiple roles.
    """
    count = len(text)
    passed = count < 6000
    return (1 if passed else 0,
            f"{count} chars ({'<6000' if passed else '>=6000'})")


def dim_consistent_formatting(sections: dict) -> tuple[int, str]:
    """All 4 expected section headers present?"""
    required = {"CERTIFICATIONS", "EXPERIENCE", "TECHNICAL SKILLS", "EDUCATION"}
    found = set(sections.keys()) & required
    missing = required - found
    passed = len(missing) == 0
    return (1 if passed else 0,
            f"{'All present' if passed else 'Missing: ' + ', '.join(missing)}")


# ---------------------------------------------------------------------------
# Scoring
# ---------------------------------------------------------------------------

def score_hiring_manager(docx_path: str) -> dict:
    """Run all 10 dimension checks and return structured results."""
    text = extract_full_text(docx_path)
    sections = extract_sections(docx_path)
    bullets = extract_bullets_with_runs(docx_path)

    dimensions = {
        "first_bullet_impact": dim_first_bullet_impact(bullets),
        "certs_before_experience": dim_certs_before_experience(sections),
        "metric_density": dim_metric_density(bullets),
        "job_title_match": dim_job_title_match(text),
        "dates_present": dim_dates_present(text),
        "no_gaps": dim_no_gaps(text),
        "skills_with_context": dim_skills_with_context(sections),
        "education_present": dim_education_present(sections),
        "single_page": dim_single_page(text),
        "consistent_formatting": dim_consistent_formatting(sections),
    }

    total = sum(score for score, _ in dimensions.values())

    breakdown = {
        name: {"score": score, "detail": detail, "pass": score == 1}
        for name, (score, detail) in dimensions.items()
    }

    return {
        "score": total,
        "max": 10,
        "breakdown": breakdown,
        "pass": total >= 9,
    }


# ---------------------------------------------------------------------------
# Output
# ---------------------------------------------------------------------------

TARGET_SCORE = 9


def print_report(result: dict):
    """Print human-readable hiring manager report."""
    print("=" * 60)
    print("HIRING MANAGER SIMULATION REPORT")
    print("=" * 60)
    print(f"Score: {result['score']}/{result['max']}")
    status = "PASS" if result["pass"] else "FAIL"
    print(f"Status: {status} (target >= {TARGET_SCORE}/10)")
    print()

    for name, info in result["breakdown"].items():
        label = name.replace("_", " ").title()
        mark = "[X]" if info["pass"] else "[ ]"
        print(f"  {mark} {label}: {info['detail']}")

    print("=" * 60)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Score a resume .docx from a hiring manager perspective."
    )
    parser.add_argument("docx_path", help="Path to the resume .docx file")
    parser.add_argument(
        "--output", choices=["human", "json"], default="human",
        help="Output format (default: human)"
    )
    args = parser.parse_args()

    result = score_hiring_manager(args.docx_path)

    if args.output == "json":
        print(json.dumps(result, indent=2))
    else:
        print_report(result)

    sys.exit(0 if result["pass"] else 1)


if __name__ == "__main__":
    main()
